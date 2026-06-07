import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SOURCE = r'D:\www.jagmagh.com\jagmagh.github.io\inputs\Solution Architect Role Re-Definition.docx'
OUTPUT = r'D:\www.jagmagh.com\jagmagh.github.io\outputs\Solution Architect Role Re-Definition v2.docx'

# ── Copy source so all styles, numbering, themes are inherited ────────────────
shutil.copy(SOURCE, OUTPUT)
doc = Document(OUTPUT)

# ── Read bullet numId from source before clearing ─────────────────────────────
src = Document(SOURCE)
bullet_numId = None
bullet_ilvl = '0'
for p in src.paragraphs:
    if p.style and p.style.name == 'List Paragraph':
        pPr = p._p.find(qn('w:pPr'))
        if pPr is not None:
            numPr = pPr.find(qn('w:numPr'))
            if numPr is not None:
                nid = numPr.find(qn('w:numId'))
                nlvl = numPr.find(qn('w:ilvl'))
                if nid is not None:
                    bullet_numId = nid.get(qn('w:val'))
                    if nlvl is not None:
                        bullet_ilvl = nlvl.get(qn('w:val'))
                    break
print(f'bullet_numId={bullet_numId}, ilvl={bullet_ilvl}')

# ── Clear body content, keep sectPr ──────────────────────────────────────────
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

# ── Colors ────────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1F, 0x38, 0x64)
GREY  = RGBColor(0x59, 0x59, 0x59)
BLUE  = RGBColor(0x2E, 0x75, 0xB6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Table cell fills
HDR_FILL    = '1F3864'
ROW_RED     = 'FCE4E4'
ROW_YELLOW  = 'FFF3D6'
ROW_GREEN   = 'E2F0E2'

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    for e in tcPr.findall(qn('w:shd')):
        tcPr.remove(e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tcPr.insert(0, shd)

def shade_para(p, hex_fill):
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn('w:shd')):
        pPr.remove(e)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    pPr.append(shd)

def add_numPr(p):
    if not bullet_numId:
        return
    pPr = p._p.get_or_add_pPr()
    for e in pPr.findall(qn('w:numPr')):
        pPr.remove(e)
    numPr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), bullet_ilvl)
    nid_el = OxmlElement('w:numId')
    nid_el.set(qn('w:val'), bullet_numId)
    numPr.append(ilvl_el)
    numPr.append(nid_el)
    pPr.insert(0, numPr)

def set_spacing(p, before=0, after=0, indent=None):
    pf = p.paragraph_format
    pf.space_before = before
    pf.space_after  = after
    if indent is not None:
        pf.left_indent = indent

# ── Content helpers ───────────────────────────────────────────────────────────

def _set_pStyle(p, style_id):
    pPr = p._p.get_or_add_pPr()
    for ex in pPr.findall(qn('w:pStyle')):
        pPr.remove(ex)
    ps = OxmlElement('w:pStyle')
    ps.set(qn('w:val'), style_id)
    pPr.insert(0, ps)

def h1(text):
    p = doc.add_paragraph()
    _set_pStyle(p, 'Heading1')
    p.add_run(text)
    return p

def h2(text):
    p = doc.add_paragraph()
    _set_pStyle(p, 'Heading2')
    p.add_run(text)
    return p

def body(text):
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=0, after=101600)
    p.add_run(text)
    return p

def sub_h(text):
    """Subsection header: 6.1, 6.2 … 7.1, 7.2 — Normal, bold navy 11.5pt."""
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=76200, after=50800)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Paragraph')
    set_spacing(p, before=0, after=63500)
    add_numPr(p)
    p.add_run(text)
    return p

def bullet_bold(bold_part, normal_part):
    p = doc.add_paragraph(style='List Paragraph')
    set_spacing(p, before=0, after=63500)
    add_numPr(p)
    r1 = p.add_run(bold_part)
    r1.bold = True
    p.add_run(normal_part)
    return p

def italic_note(bold_label, rest, after=101600):
    """Carries into delivery / Also applies / AI advisory — indent, bold+italic grey 10pt."""
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=38100, after=after, indent=228600)
    r1 = p.add_run(bold_label)
    r1.bold = True; r1.italic = True
    r1.font.size = Pt(10); r1.font.color.rgb = GREY
    r2 = p.add_run(rest)
    r2.italic = True
    r2.font.size = Pt(10); r2.font.color.rgb = GREY
    return p

def how_header():
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=50800, after=25400)
    r = p.add_run('How and when this is engaged')
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = GREY
    return p

def how_line(label, rest):
    """Trigger / Engagement shape / Staffing — indented, blue bold label 9.5pt."""
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=0, after=25400, indent=228600)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(9.5); r1.font.color.rgb = BLUE
    r2 = p.add_run(rest)
    r2.font.size = Pt(9.5)
    return p

def lifecycle_line(label, rest):
    """Project setup / Throughout delivery / Key gates — indented, blue bold 10pt."""
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=0, after=50800, indent=228600)
    r1 = p.add_run(label)
    r1.bold = True; r1.font.size = Pt(10); r1.font.color.rgb = BLUE
    r2 = p.add_run(rest)
    r2.font.size = Pt(10)
    return p

def engagement_header(text):
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=50800, after=25400)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY
    return p

def filter_q(text):
    """Governing filter question — callout: thick blue left border, light blue bg, italic navy 12pt."""
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=76200, after=76200)
    pf = p.paragraph_format
    pf.left_indent = 228600
    pf.right_indent = 228600
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side, sz, space, color in [
        ('top',    '2',  '8',  'D9D9D9'),
        ('left',   '18', '12', '2E75B6'),
        ('bottom', '2',  '8',  'D9D9D9'),
        ('right',  '2',  '8',  'D9D9D9'),
    ]:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), space)
        b.set(qn('w:color'), color)
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F4F8FB')
    pPr.append(shd)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    return p

def end_portrait_section(p):
    """Inline sectPr ending a portrait US Letter section at paragraph p."""
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    t = OxmlElement('w:type'); t.set(qn('w:val'), 'nextPage'); sectPr.append(t)
    sz = OxmlElement('w:pgSz'); sz.set(qn('w:w'), '12240'); sz.set(qn('w:h'), '15840'); sectPr.append(sz)
    m = OxmlElement('w:pgMar')
    m.set(qn('w:left'),'1440'); m.set(qn('w:right'),'1440')
    m.set(qn('w:top'),'1440');  m.set(qn('w:bottom'),'1440')
    m.set(qn('w:header'),'708'); m.set(qn('w:footer'),'708'); m.set(qn('w:gutter'),'0')
    sectPr.append(m)
    pPr.append(sectPr)

def end_landscape_section(p):
    """Inline sectPr ending a landscape US Letter section at paragraph p."""
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    t = OxmlElement('w:type'); t.set(qn('w:val'), 'nextPage'); sectPr.append(t)
    sz = OxmlElement('w:pgSz')
    sz.set(qn('w:orient'),'landscape'); sz.set(qn('w:w'), '15840'); sz.set(qn('w:h'), '12240')
    sectPr.append(sz)
    m = OxmlElement('w:pgMar')
    m.set(qn('w:left'),'1080'); m.set(qn('w:right'),'1080')
    m.set(qn('w:top'),'1080');  m.set(qn('w:bottom'),'1080')
    m.set(qn('w:header'),'708'); m.set(qn('w:footer'),'708'); m.set(qn('w:gutter'),'0')
    sectPr.append(m)
    pPr.append(sectPr)

def set_cell_text(cell, text, bold=False, color=BLACK, size=Pt(9)):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.bold = bold; r.font.size = size; r.font.color.rgb = color

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE + SUBTITLE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph(style='Normal')
set_spacing(p, before=152400, after=38100)
r = p.add_run('Solution Architect Role Definition')
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = NAVY

p = doc.add_paragraph(style='Normal')
set_spacing(p, before=0, after=152400)
r = p.add_run('The Architect’s Operating Model in an AI-Native World')
r.italic = True; r.font.size = Pt(13); r.font.color.rgb = GREY

# ══════════════════════════════════════════════════════════════════════════════
#  1. PURPOSE
# ══════════════════════════════════════════════════════════════════════════════
h1('1. Purpose')
body('This document defines the architect’s role in an environment where generative AI has materially changed how technical work is produced. Its purpose is to establish where the architect’s value now concentrates, how that value is delivered to clients and project teams, and what the architect must do internally to keep that value credible and current.')
body('The core premise is that AI has compressed the cost and time of construction and the recall of technical breadth. As a result, the architect’s value shifts from knowing and building toward judging, directing, and being accountable for technical decisions. The role is structured accordingly: the client-facing work is defined in the sections that follow, grounded in an internal practice that keeps it credible.')

# ══════════════════════════════════════════════════════════════════════════════
#  2. OPERATING PREMISE
# ══════════════════════════════════════════════════════════════════════════════
h1('2. Operating Premise')
body('Three shifts in the technical landscape define this role:')
bullet_bold('Construction is commoditised. ', 'AI, applied properly, produces working implementations faster and at lower cost than traditional hand-construction. Building as a standalone billable activity is increasingly difficult to justify as work requiring an architect.')
bullet_bold('Breadth and recall are commoditised. ', 'The encyclopedic knowledge of technologies, products, and patterns that once set senior architects apart is now available on demand. The premium attached to simply knowing more has eroded.')
bullet_bold('Judgment has not been commoditised. ', 'AI is a strong generator and a weak discriminator: it produces plausible options confidently but carries no accountability for whether they survive contact with reality. Deciding what to build, what not to build, and what will fail later remains scarce, and this is where the experience and judgment of the architect can still add value.')
body('The role is therefore designed so that judgment is the primary offering, and construction is the internal discipline that keeps it grounded.')

# ══════════════════════════════════════════════════════════════════════════════
#  3. WHAT JUDGMENT IS  (new section)
# ══════════════════════════════════════════════════════════════════════════════
h1('3. What Judgment Is')
body('Judgment is the term used throughout this document to describe what the architect offers where AI is structurally weakest. It is worth naming concretely, because it is easy to invoke and hard to specify.')
body('In practice, judgment means:')
bullet('Identifying requirements that are traps before accepting them — where the stated need diverges from the actual problem the client needs to solve.')
bullet('Distinguishing a solution that demonstrates well in a pitch from one that will survive delivery at scale and under real operating conditions.')
bullet('Reading the gap between what a client says they want and what their organisation actually needs, can absorb, and is ready to change.')
bullet('Knowing when the technically correct recommendation is politically unrealizable, and navigating to one that can actually be acted on.')
bullet('Pattern-matching failure modes from prior engagements onto a new situation before construction begins, not after it has gone wrong.')
bullet('Deciding what risk is knowingly acceptable for this client, in this context — rather than eliminating all risk indiscriminately or accepting it without naming it.')
bullet('Knowing when the honest answer is “don’t build this” — and having the standing and the relationship to say it and be heard.')
body('None of these are computable from a prompt. They require accumulated exposure to how engagements actually fail, not just what they look like when they are going well.')

# ══════════════════════════════════════════════════════════════════════════════
#  4. HOW THE ROLE HAS CHANGED
# ══════════════════════════════════════════════════════════════════════════════
h1("4. How the Role Has Changed")
p_last_portrait = body("The table below compares the architect's activities before and after the widespread adoption of AI, and states why each has shifted. The pattern is consistent: where AI now produces the output or supplies the recall, the activity no longer requires an architect and can be delivered at developer level; where judgment, context, and accountability are required, the activity remains architect-level work and in several cases becomes more valuable.")
end_portrait_section(p_last_portrait)

# Legend title
p = doc.add_paragraph(style='Normal')
set_spacing(p, before=76200, after=25400)
r = p.add_run('Before AI vs. Now — and Why')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY

# Color legend rows (shaded paragraphs)
for fill, txt in [
    (ROW_RED,    '   No longer needs an architect — can be delivered at developer level'),
    (ROW_YELLOW, '   Production drops to developer level; architect directs and validates'),
    (ROW_GREEN,  '   Still the architect’s — unaffected or more valuable'),
]:
    p = doc.add_paragraph(style='Normal')
    set_spacing(p, before=0, after=0)
    r = p.add_run(txt)
    r.font.size = Pt(9)
    shade_para(p, fill)

# Gap before table
p = doc.add_paragraph(style='Normal')
set_spacing(p, before=0, after=76200)

# Main table
tbl = doc.add_table(rows=9, cols=6)

# Table borders — thin white lines between cells
def set_table_borders(table):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'FFFFFF')
        borders.append(b)
    for ex in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(ex)
    tblPr.append(borders)

set_table_borders(tbl)

# Pin table to full landscape available width (13680 twips = 9.5" - 0.75" margins each side)
def set_table_width(table, twips):
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr)
    for ex in tblPr.findall(qn('w:tblW')): tblPr.remove(ex)
    w = OxmlElement('w:tblW')
    w.set(qn('w:w'), str(twips)); w.set(qn('w:type'), 'dxa')
    tblPr.append(w)
    for ex in tblPr.findall(qn('w:tblLayout')): tblPr.remove(ex)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)

set_table_width(tbl, 13680)

# Column widths scaled to landscape available width (~9.2" total)
col_w = [Inches(1.8), Inches(1.1), Inches(1.65), Inches(1.65), Inches(1.65), Inches(1.35)]
for row in tbl.rows:
    for ci, cell in enumerate(row.cells):
        cell.width = col_w[ci]

# Header row
for ci, hdr in enumerate(['Activity','Lifecycle Stage','Role Before AI','Role Now','Why It Shifted','Mapped Offering (Sec. 6)']):
    shade_cell(tbl.rows[0].cells[ci], HDR_FILL)
    set_cell_text(tbl.rows[0].cells[ci], hdr, bold=True, color=WHITE, size=Pt(10.5))

# Data rows: (fill, [(text, bold, color), ...])
data_rows = [
  (ROW_RED, [
    ('Technical breadth — knowing the landscape',True,NAVY),
    ('Cross-cutting',False,GREY),
    ('Encyclopedic recall of technologies, products, and patterns was a core differentiator.',False,BLACK),
    ('AI supplies breadth on demand; the architect curates and selects rather than recalls.',False,BLACK),
    ('Recall is now instant and free, so it no longer justifies an architect and can sit with developers.',False,BLACK),
    ('Not sold — developer / AI',True,NAVY),
  ]),
  (ROW_RED, [
    ('Building and construction',True,NAVY),
    ('Build',False,GREY),
    ('Significant hands-on construction was central and billable.',False,BLACK),
    ('Construction is delegated to AI; building is retained internally only to stay grounded.',False,BLACK),
    ('AI builds faster and cheaper, so construction can be delivered at developer level rather than by an architect.',False,BLACK),
    ('Not sold — developer / AI (retained internally per Sec. 6)',True,NAVY),
  ]),
  (ROW_YELLOW, [
    ('Producing solutions and proposals',True,NAVY),
    ('Pre-sales',False,GREY),
    ('Hand-crafting solution outlines, proposals, estimates, and demos.',False,BLACK),
    ('Directing and validating AI-produced artifacts; owning qualification and the client relationship.',False,BLACK),
    ('AI produces the artifacts quickly; value moves to judgment, feasibility, and trust.',False,BLACK),
    ('Sold as 5.1 — architect judgment guides AI / Dev output',True,NAVY),
  ]),
  (ROW_YELLOW, [
    ('Design patterns and first-draft deliverables',True,NAVY),
    ('Design',False,GREY),
    ('Producing reference designs, matrices, and first-draft documents from expertise.',False,BLACK),
    ('AI generates the first draft; the architect judges, corrects, and adapts to context.',False,BLACK),
    ('Generation is commoditised; discrimination and fit-to-context are not.',False,BLACK),
    ('Feeds 5.1 & 5.5 — architect judges AI drafts',True,NAVY),
  ]),
  (ROW_YELLOW, [
    ('Reviewing and assuring the build',True,NAVY),
    ('Build',False,GREY),
    ('Reviewing deliverables produced by human teams.',False,BLACK),
    ('First-pass review by AI and developers; the architect owns the trade-offs and risk the review surfaces.',False,BLACK),
    ('Detection is commoditised, but adjudicating conflicting goals and accepting risk is not.',False,BLACK),
    ('Sold as 5.5 — architect owns trade-offs & risk',True,NAVY),
  ]),
  (ROW_GREEN, [
    ('Deciding what to build vs. what not to build',True,NAVY),
    ('Pre-sales / Design',False,GREY),
    ('An implicit part of architect judgment, not always named as a distinct offering.',False,BLACK),
    ('An explicit, high-value offering — the scarce discriminator over confident AI output.',False,BLACK),
    ('AI generates options confidently but cannot decide what is worth doing; this rises in value.',False,BLACK),
    ('Sold as 5.2 — irreducibly architect-led',True,NAVY),
  ]),
  (ROW_GREEN, [
    ('Platform, vendor, and transformation decisions',True,NAVY),
    ('Design / Plan',False,GREY),
    ('Assessments and selections grounded in experience and client context.',False,BLACK),
    ('Largely unchanged; AI assists analysis but the decision and accountability remain human.',False,BLACK),
    ('Requires context and accountability AI cannot hold, so it remains architect-level work.',False,BLACK),
    ('Sold as 5.3 — irreducibly architect-led',True,NAVY),
  ]),
  (ROW_GREEN, [
    ('Client relationship and accountability',True,NAVY),
    ('Cross-cutting',False,GREY),
    ('The trusted relationship and the willingness to stake credibility on a recommendation.',False,BLACK),
    ('Unchanged and more pivotal as AI-polished competitors look increasingly alike.',False,BLACK),
    ('Trust, presence, and accountability are inherently human and cannot be automated.',False,BLACK),
    ('Underpins 5.1 & 5.4 — the connective tissue',True,NAVY),
  ]),
]

for ri, (fill, cells) in enumerate(data_rows):
    for ci, (txt, bold, color) in enumerate(cells):
        shade_cell(tbl.rows[ri+1].cells[ci], fill)
        set_cell_text(tbl.rows[ri+1].cells[ci], txt, bold=bold, color=color, size=Pt(9))

# Empty paragraph ending the landscape section (returns to portrait)
p_end_landscape = doc.add_paragraph(style='Normal')
end_landscape_section(p_end_landscape)

# ══════════════════════════════════════════════════════════════════════════════
#  5. WHERE THE ARCHITECT’S VALUE HAS GRAVITATED
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Where the Architect’s Value Has Gravitated")
body('Each offering in this section is an expression of senior technical judgment applied where decisions carry the most cost and where AI is structurally weakest. The offerings divide into two categories by how they are engaged: independent sellable engagements that stand on their own, and roles performed within a larger implementation effort. Most offerings have a primary category, with a note on where they also apply to the other.')

h2('5A. Independent Sellable Engagements')
body('These can be scoped, sold, and delivered as engagements in their own right, independent of any single implementation project.')

# 5.1
sub_h('5.1  Pre-Sales Solutioning')
body('What the client gets is an architect who shapes and vouches for the proposed solution — not just produces it. In a market where AI-generated proposals are uniformly polished, the differentiator is the judgment behind them: whether the solution is real, whether the estimate is honest, and whether the people proposing it have the standing to deliver.')
bullet('Making the call on which opportunities are worth pursuing and which should be declined.')
bullet('Shaping and validating the solution, estimate, and approach — directing AI and developer output rather than producing it by hand.')
bullet('Vouching for feasibility: confirming the proposed approach will survive delivery and being explicit about what is not yet production-ready.')
bullet('Staking personal credibility on the recommendation in the room, not just the document.')
italic_note('Carries into delivery:  ', 'the trust established here is the foundation the implementation role (5.5) is built on.')
how_header()
how_line('Trigger — ', 'An active or prospective opportunity where a solution and proposal must be shaped to win the work.')
how_line('Engagement shape — ', 'Bid and pursuit work — typically unbilled investment ahead of a sale, scoped to the opportunity.')
how_line('Staffing — ', 'Architect-led on qualification, feasibility, and the client relationship; proposal, estimate, and demo production carried by developers and AI under the architect’s direction.')

# 5.2
sub_h('5.2  Decision Guidance — What to Build vs. What Not to Build')
body('The highest-value protection a client can buy: an architect who will tell them not to build something before they commit the budget to build it wrong. Incorrect decisions at this layer are the most expensive, and AI — which generates options confidently but carries no accountability for their consequences — is least reliable precisely here.')
bullet('Advising which initiatives are worth pursuing, which should be deferred, and which should be stopped.')
bullet('Identifying requirements that are traps and challenging solutions that are unnecessary or misdirected before they consume delivery capacity.')
bullet('Making the call on where AI genuinely belongs in the solution and where a conventional, deterministic approach is the right answer.')
bullet('Framing the trade-offs clearly so that stakeholders can commit to a decision rather than defer one.')
italic_note('Also applies within implementations:  ', 'as a decision checkpoint when scope, requirements, or direction change mid-project.')
how_header()
how_line('Trigger — ', 'A client facing a consequential, contested, or expensive decision about what to pursue.')
how_line('Engagement shape — ', 'A short, sharp advisory engagement — an assessment, decision review, or workshop — or a named decision-owner role within a larger programme.')
how_line('Staffing — ', 'Architect-led throughout; AI used to surface and stress-test options, with the architect owning the recommendation.')

# 5.3
sub_h('5.3  Platform, Vendor, and Transformation Decisions')
body('What the client gets is a recommendation grounded in their actual situation — not a feature matrix or a vendor pitch. Platform and transformation decisions require contextual and relational knowledge that AI cannot access: the organisation’s operating constraints, its political dynamics, what its teams can realistically absorb, and what prior decisions have foreclosed. These remain durable architect responsibilities regardless of how capable tooling becomes.')
bullet('Assessing the current landscape honestly — including constraints and technical debt the client may prefer not to name.')
bullet('Selecting platforms and vendors against fit, risk, and total cost rather than marketed capability.')
bullet('Sequencing transformation against organisational readiness, not against an idealised roadmap.')
bullet('Threading recommendations through how the client actually operates, not how it should theoretically operate.')
italic_note('Also applies within implementations:  ', 'when a platform or vendor decision is reopened during delivery due to new constraints.')
how_header()
how_line('Trigger — ', 'A client planning a transformation, replatforming, or major procurement, or needing an independent view of their landscape.')
how_line('Engagement shape — ', 'A discrete, scoped engagement — assessment, vendor or platform selection, or transformation roadmap — that typically precedes and seeds the programme that follows.')
how_line('Staffing — ', 'Architect-led; AI accelerates landscape analysis and option comparison, with the architect owning the assessment conclusions and the recommendation.')

# 5.4
sub_h('5.4  Trusted Technology Advisory')
body('What the client gets is an independent, credible voice that will tell leadership what it does not want to hear — including recommending against a technology, a vendor, or an initiative when the honest assessment is that it is the wrong choice. The value is not technology-specific; it is the combination of judgment, context, and willingness to stake reputation on a recommendation that runs contrary to the prevailing preference in the room.')
bullet('Helping leadership distinguish substance from hype and calibrate realistic expectations — particularly under the pressure of technology cycles.')
bullet('Advising on risk, security, and governance posture, and on whether the organisation is actually ready for a given change.')
bullet('Guiding technology investment and its sequencing against what the organisation can absorb and sustain.')
bullet('Recommending against a technology — including AI — where its limitations, cost, or organisational readiness make it the wrong fit.')
italic_note('AI advisory as a specialisation:  ', 'AI-specific advisory is carved out as a distinct offering within this — defining where AI genuinely belongs in a client’s architecture, setting the governance and guardrails, and being equally clear where the honest answer is that AI is not the right fit.', after=38100)
italic_note('Also applies within implementations:  ', 'as the governance and standards authority that build work is held against.')
how_header()
how_line('Trigger — ', 'Leadership needing a credible, independent technology voice — often continuously rather than for a single decision.')
how_line('Engagement shape — ', 'A relationship rather than a project — retainer-style, an enterprise-architecture seat, or EA-team augmentation, sold on reputation and continuity.')
how_line('Staffing — ', 'Architect-only at the advisory seat; analysis and supporting material prepared by developers and AI, with the architect owning the counsel given.')

h2('5B. Roles Within a Larger Implementation')
body('This offering exists only when a build is underway. It is engaged as a defined role inside an implementation project rather than sold as a standalone engagement.')

# 5.5
sub_h('5.5  Architectural Decisions and Risk Ownership during the Build')
body('First-pass review and assessment are increasingly done by developers and AI. The architect’s value in delivery is not running those checks but owning the trade-off decisions they surface and being accountable for the risk that ships.')
bullet('Owning the architectural trade-offs — reliability, performance, cost, security, and speed-to-market are competing goods, and the right balance depends on this client’s context and risk appetite.')
bullet('Adjudicating the findings of AI- and developer-led reviews and assurance, deciding what must be fixed, what is acceptable, and what risk is knowingly accepted.')
bullet('Setting the architecture, standards, and guardrails up front that the build is then held against.')
bullet('Standing accountable for what is released — the decision and the risk acceptance, not the detection.')
engagement_header('Engagement across the implementation lifecycle')
body('The role is heaviest at the start and at key decision gates, not uniformly throughout.')
lifecycle_line('Project setup — ', 'high involvement — establish the architecture, standards, guardrails, and development approach the build will follow.')
lifecycle_line('Throughout delivery — ', 'light, advisory involvement — available for design questions and to assess the impact of changing requirements.')
lifecycle_line('Key gates — ', 'concentrated involvement — own the trade-off and risk decisions on major deliverables and design changes before they are committed.')
italic_note('Also applies as a standalone engagement:  ', 'as an independent architecture or assurance review commissioned outside a delivery role.')
how_header()
how_line('Trigger — ', 'An implementation under way where consequential architectural trade-offs and release-risk decisions must be owned.')
how_line('Engagement shape — ', 'A defined role within a delivery programme — front-loaded at setup, then concentrated at decision gates rather than continuous.')
how_line('Staffing — ', 'Architect owns decisions and risk acceptance; first-pass review, assurance, and construction carried by developers and AI.')

# ══════════════════════════════════════════════════════════════════════════════
#  6. SUPPORTING DISCIPLINE: BUILDING FOR KNOWLEDGE
# ══════════════════════════════════════════════════════════════════════════════
h1('6. Supporting Discipline: Building for Knowledge')
body('The advisory work described in the preceding sections is only as credible as the judgment behind it. That judgment is kept honest by a deliberate, selective building practice — one that is not sold to clients and is not designed to compete on construction speed or volume, but is maintained because advisory credibility that is not tested against technical reality eventually drifts from it.')

sub_h('6.1  Activities')
bullet('Building working prototypes and proofs-of-concept in the AI-native stack — agentic workflows, retrieval-augmented patterns, evaluation strategies, and model orchestration — to establish first-hand what performs versus what only demonstrates well.')
bullet('Maintaining sufficient currency with cloud AI services (e.g. AWS Bedrock, Azure AI Foundry) and the underlying data platform plumbing to distinguish substance from vendor or team hand-waving.')
bullet('Building selectively before advising, so that recommendations on a pattern are informed by direct exposure to its failure modes and sharp edges.')

sub_h('6.2  Operating Principles')
bullet('This work is judged by whether it keeps the architect close to ground truth, not by whether it is billable.')
bullet('Specific tooling knowledge has a short half-life; the objective is grounding and the ability to detect what is real, not exhaustive mastery of current tools.')
bullet('If hands-on practice lapses, advisory judgment silently drifts from reality. Maintaining this practice is therefore a non-negotiable input to the credibility of the advisory work.')

# ══════════════════════════════════════════════════════════════════════════════
#  7. THE GOVERNING FILTER
# ══════════════════════════════════════════════════════════════════════════════
h1('7. The Governing Filter')
body('Because both construction and advisory work can be performed competently, the deciding question for any engagement is not which type of work it is, but whether it compounds:')
filter_q('Does performing this work well make the architect harder to replace and more trusted next year — or does it simply produce one more competent output that anyone with current tooling could have produced?')
body('Internal building passes this filter only when it demonstrably sharpens the advisory and assurance work. External engagements are prioritised where they make the architect harder to replace and more trusted, rather than where they merely produce volume.')

# ══════════════════════════════════════════════════════════════════════════════
#  8. SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1('8. Summary')
body('The architect in an AI-native world places judgment, decision guidance, and accountability on the invoice, while retaining a deliberate, selective building practice as the credibility floor that keeps that judgment honest. Construction is no longer the headline offering; it is the root system beneath it. The role is defined not by a choice between building and advising, but by an operating model in which directed, accountable judgment is the value, and hands-on practice is the discipline that keeps the value real.')

doc.save(OUTPUT)
print('Done:', OUTPUT)
