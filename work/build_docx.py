from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'D:\www.jagmagh.com\jagmagh.github.io\inputs\Solution Architect Role Re-Definition v2.docx'

doc = Document()

# Page margins
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(3.0)

# Style helpers
def h1(text):
    p = doc.add_paragraph()
    p.style = 'Heading 1'
    p.add_run(text)

def h2(text):
    p = doc.add_paragraph()
    p.style = 'Heading 2'
    p.add_run(text)

def body(text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.add_run(text)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)

def italic_note(text):
    p = doc.add_paragraph()
    p.style = 'Normal'
    r = p.add_run(text)
    r.italic = True

def bold_label(label, rest=''):
    p = doc.add_paragraph()
    p.style = 'Normal'
    r1 = p.add_run(label)
    r1.bold = True
    if rest:
        p.add_run(rest)

# ─── TITLE ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Solution Architect Role Definition')
r.bold = True
r.font.size = Pt(22)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('The Architect’s Operating Model in an AI-Native World')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ─── 1. PURPOSE ───────────────────────────────────────────────────────────────
h1('1. Purpose')
body('This document defines the role of the Solution Architect in an environment where generative AI has materially changed how technical work is produced. Its purpose is to establish where the architect’s value now concentrates, how that value is delivered to clients and project teams, and what the architect must do internally to keep that value credible and current.')
body('The core premise is that AI has compressed the cost and time of construction and the recall of technical breadth. As a result, the architect’s value shifts from knowing and building toward judging, directing, and being accountable for technical decisions. The role is structured accordingly into two layers: an internal capability layer that keeps the architect grounded, and an external value layer that defines what the organisation offers and bills for.')

# ─── 2. OPERATING PREMISE ────────────────────────────────────────────────────
h1('2. Operating Premise')
body('Three shifts in the technical landscape define this role:')
bullet('Construction is commoditised. AI, applied properly, produces working implementations faster and at lower cost than traditional hand-construction. Building as a standalone billable activity is increasingly difficult to justify as work requiring an architect.')
bullet('Breadth and recall are commoditised. The encyclopedic knowledge of technologies, products, and patterns that once set senior architects apart is now available on demand. The premium attached to simply knowing more has eroded.')
bullet('Judgment has not been commoditised. AI is a strong generator and a weak discriminator: it produces plausible options confidently but carries no accountability for whether they survive contact with reality. Deciding what to build, what not to build, and what will fail later remains scarce, and this is where the experience and judgment of the architect can still add value.')
body('The role is therefore designed so that judgment is the primary offering, construction is retained as an internal grounding discipline, and the two reinforce each other.')

# ─── 3. WHAT JUDGMENT IS (NEW) ───────────────────────────────────────────────
h1('3. What Judgment Is')
body('Judgment is the term used throughout this document to describe what the architect offers where AI is structurally weakest. It is worth naming concretely, because it is easy to invoke and hard to specify.')
body('In practice, judgment means:')
bullet('Identifying requirements that are traps before accepting them — where the stated need diverges from the actual problem the client needs to solve.')
bullet('Distinguishing a solution that demonstrates well in a pitch from one that will survive delivery at scale and under real operating conditions.')
bullet('Reading the gap between what a client says they want and what their organisation actually needs, can absorb, and is ready to change.')
bullet('Knowing when the technically correct recommendation is politically unrealizable, and navigating to one that can actually be acted on.')
bullet('Pattern-matching failure modes from prior engagements onto a new situation before construction begins, not after it has gone wrong.')
bullet('Deciding what risk is knowingly acceptable for this client, in this context — rather than eliminating all risk indiscriminately or accepting it without naming it.')
bullet('Knowing when the honest answer is “don’t build this” — and having the standing and the relationship to say it and be heard.')
body('None of these are computable from a prompt. They require accumulated exposure to how engagements actually fail, not just what they look like when they are going well.')

# ─── 4. HOW THE ROLE HAS CHANGED ─────────────────────────────────────────────
h1('4. How the Role Has Changed')
body('The table below compares the architect’s activities before and after the widespread adoption of AI, and states why each has shifted. The pattern is consistent: where AI now produces the output or supplies the recall, the activity no longer requires an architect and can be delivered at developer level; where judgment, context, and accountability are required, the activity remains architect-level work and in several cases becomes more valuable.')

bold_label('Before AI vs. Now — and Why')
italic_note('■  No longer needs an architect — can be delivered at developer level')
italic_note('■  Production drops to developer level; architect directs and validates')
italic_note('■  Still the architect’s — unaffected or more valuable')

# Table
tbl = doc.add_table(rows=9, cols=6)
tbl.style = 'Table Grid'

# Set column widths
col_widths = [Cm(3.8), Cm(2.4), Cm(3.8), Cm(3.8), Cm(3.8), Cm(3.4)]
for row in tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = col_widths[i]

headers = ['Activity', 'Lifecycle Stage', 'Role Before AI', 'Role Now', 'Why It Shifted', 'Mapped Offering\n(Sec. 6)']
for i, hdr in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(hdr)
    r.bold = True
    r.font.size = Pt(9)

rows_data = [
    ['Technical breadth — knowing the landscape', 'Cross-cutting',
     'Encyclopedic recall of technologies, products, and patterns was a core differentiator.',
     'AI supplies breadth on demand; the architect curates and selects rather than recalls.',
     'Recall is now instant and free, so it no longer justifies an architect.',
     'Not sold — developer / AI'],
    ['Building and construction', 'Build',
     'Significant hands-on construction was central and billable.',
     'Construction is delegated to AI; building is retained internally only to stay grounded.',
     'AI builds faster and cheaper, so construction can be delivered at developer level.',
     'Not sold — developer / AI\n(retained internally per Sec. 7)'],
    ['Producing solutions and proposals', 'Pre-sales',
     'Hand-crafting solution outlines, proposals, estimates, and demos.',
     'Directing and validating AI-produced artifacts; owning qualification and the client relationship.',
     'AI produces the artifacts quickly; value moves to judgment, feasibility, and trust.',
     'Sold as 6.1 — architect judgment guides AI / Dev output'],
    ['Design patterns and first-draft deliverables', 'Design',
     'Producing reference designs, matrices, and first-draft documents from expertise.',
     'AI generates the first draft; the architect judges, corrects, and adapts to context.',
     'Generation is commoditised; discrimination and fit-to-context are not.',
     'Feeds 6.1 & 6.5 — architect judges AI drafts'],
    ['Reviewing and assuring the build', 'Build',
     'Reviewing deliverables produced by human teams.',
     'First-pass review by AI and developers; the architect owns the trade-offs and risk the review surfaces.',
     'Detection is commoditised, but adjudicating conflicting goals and accepting risk is not.',
     'Sold as 6.5 — architect owns trade-offs & risk'],
    ['Deciding what to build vs. what not to build', 'Pre-sales / Design',
     'An implicit part of architect judgment, not always named as a distinct offering.',
     'An explicit, high-value offering — the scarce discriminator over confident AI output.',
     'AI generates options confidently but cannot decide what is worth doing; this rises in value.',
     'Sold as 6.2 — irreducibly architect-led'],
    ['Platform, vendor, and transformation decisions', 'Design / Plan',
     'Assessments and selections grounded in experience and client context.',
     'Largely unchanged; AI assists analysis but the decision and accountability remain human.',
     'Requires context and accountability AI cannot hold, so it remains architect-level work.',
     'Sold as 6.3 — irreducibly architect-led'],
    ['Client relationship and accountability', 'Cross-cutting',
     'The trusted relationship and the willingness to stake credibility on a recommendation.',
     'Unchanged and more pivotal as AI-polished competitors look increasingly alike.',
     'Trust, presence, and accountability are inherently human and cannot be automated.',
     'Underpins 6.1 & 6.4 — the connective tissue'],
]

for ri, row_data in enumerate(rows_data):
    row = tbl.rows[ri + 1]
    for ci, text in enumerate(row_data):
        cell = row.cells[ci]
        cell.text = ''
        p = cell.paragraphs[0]
        p.add_run(text).font.size = Pt(9)

# ─── 5. WHAT THE ARCHITECT OFFERS ────────────────────────────────────────────
h1('5. What the Architect Offers')
body('The architect’s work divides into two parts. The first is an internal building practice — selective, deliberate, and not billed to clients — whose purpose is to keep the architect’s judgment grounded in technical reality rather than drifting toward theory. The second is the advisory and assurance work offered to clients: directing, deciding, and standing accountable for technical outcomes. Only the second is positioned and sold; the first is what keeps it credible. The offerings that follow define the external layer.')

# ─── 6. EXTERNAL VALUE LAYER ─────────────────────────────────────────────────
h1('6. External Value Layer — The Billable Surface')
body('This layer defines what the architect offers to clients and project teams. Each offering is an expression of senior technical judgment applied where decisions carry the most cost and where AI is structurally weakest. The offerings divide into two categories by how they are engaged: independent sellable engagements that stand on their own, and roles performed within a larger implementation effort. Most offerings have a primary category, with a note on where they also apply to the other.')

h2('6A.  Independent Sellable Engagements')
body('These can be scoped, sold, and delivered as engagements in their own right, independent of any single implementation project.')

h2('6.1  Pre-Sales Solutioning')
body('What the client gets is an architect who shapes and vouches for the proposed solution — not just produces it. In a market where AI-generated proposals are uniformly polished, the differentiator is the judgment behind them: whether the solution is real, whether the estimate is honest, and whether the people proposing it have the standing to deliver.')
bullet('Making the call on which opportunities are worth pursuing and which should be declined.')
bullet('Shaping and validating the solution, estimate, and approach — directing AI and developer output rather than producing it by hand.')
bullet('Vouching for feasibility: confirming the proposed approach will survive delivery and being explicit about what is not yet production-ready.')
bullet('Staking personal credibility on the recommendation in the room, not just the document.')
italic_note('Carries into delivery:  the trust established here is the foundation the implementation role (6.5) is built on.')
bold_label('How and when this is engaged')
bold_label('Trigger — ', 'An active or prospective opportunity where a solution and proposal must be shaped to win the work.')
bold_label('Engagement shape — ', 'Bid and pursuit work — typically unbilled investment ahead of a sale, scoped to the opportunity.')
bold_label('Staffing — ', 'Architect-led on qualification, feasibility, and the client relationship; proposal, estimate, and demo production carried by developers and AI under the architect’s direction.')

h2('6.2  Decision Guidance — What to Build vs. What Not to Build')
body('The highest-value protection a client can buy: an architect who will tell them not to build something before they commit the budget to build it wrong. Incorrect decisions at this layer are the most expensive, and AI — which generates options confidently but carries no accountability for their consequences — is least reliable precisely here.')
bullet('Advising which initiatives are worth pursuing, which should be deferred, and which should be stopped.')
bullet('Identifying requirements that are traps and challenging solutions that are unnecessary or misdirected before they consume delivery capacity.')
bullet('Making the call on where AI genuinely belongs in the solution and where a conventional, deterministic approach is the right answer.')
bullet('Framing the trade-offs clearly so that stakeholders can commit to a decision rather than defer one.')
italic_note('Also applies within implementations:  as a decision checkpoint when scope, requirements, or direction change mid-project.')
bold_label('How and when this is engaged')
bold_label('Trigger — ', 'A client facing a consequential, contested, or expensive decision about what to pursue.')
bold_label('Engagement shape — ', 'A short, sharp advisory engagement — an assessment, decision review, or workshop — or a named decision-owner role within a larger programme.')
bold_label('Staffing — ', 'Architect-led throughout; AI used to surface and stress-test options, with the architect owning the recommendation.')

h2('6.3  Platform, Vendor, and Transformation Decisions')
body('What the client gets is a recommendation grounded in their actual situation — not a feature matrix or a vendor pitch. Platform and transformation decisions require contextual and relational knowledge that AI cannot access: the organisation’s operating constraints, its political dynamics, what its teams can realistically absorb, and what prior decisions have foreclosed. These remain durable architect responsibilities regardless of how capable tooling becomes.')
bullet('Assessing the current landscape honestly — including constraints and technical debt the client may prefer not to name.')
bullet('Selecting platforms and vendors against fit, risk, and total cost rather than marketed capability.')
bullet('Sequencing transformation against organisational readiness, not against an idealised roadmap.')
bullet('Threading recommendations through how the client actually operates, not how it should theoretically operate.')
italic_note('Also applies within implementations:  when a platform or vendor decision is reopened during delivery due to new constraints.')
bold_label('How and when this is engaged')
bold_label('Trigger — ', 'A client planning a transformation, replatforming, or major procurement, or needing an independent view of their landscape.')
bold_label('Engagement shape — ', 'A discrete, scoped engagement — assessment, vendor or platform selection, or transformation roadmap — that typically precedes and seeds the programme that follows.')
bold_label('Staffing — ', 'Architect-led; AI accelerates landscape analysis and option comparison, with the architect owning the assessment conclusions and the recommendation.')

h2('6.4  Trusted Technology Advisory')
body('What the client gets is an independent, credible voice that will tell leadership what it does not want to hear — including recommending against a technology, a vendor, or an initiative when the honest assessment is that it is the wrong choice. The value is not technology-specific; it is the combination of judgment, context, and willingness to stake reputation on a recommendation that runs contrary to the prevailing preference in the room.')
bullet('Helping leadership distinguish substance from hype and calibrate realistic expectations — particularly under the pressure of technology cycles.')
bullet('Advising on risk, security, and governance posture, and on whether the organisation is actually ready for a given change.')
bullet('Guiding technology investment and its sequencing against what the organisation can absorb and sustain.')
bullet('Recommending against a technology — including AI — where its limitations, cost, or organisational readiness make it the wrong fit.')
italic_note('AI advisory as a specialisation:  AI-specific advisory is carved out as a distinct offering within this — defining where AI genuinely belongs in a client’s architecture, setting the governance and guardrails, and being equally clear where the honest answer is that AI is not the right fit.')
italic_note('Also applies within implementations:  as the governance and standards authority that build work is held against.')
bold_label('How and when this is engaged')
bold_label('Trigger — ', 'Leadership needing a credible, independent technology voice — often continuously rather than for a single decision.')
bold_label('Engagement shape — ', 'A relationship rather than a project — retainer-style, an enterprise-architecture seat, or EA-team augmentation, sold on reputation and continuity.')
bold_label('Staffing — ', 'Architect-only at the advisory seat; analysis and supporting material prepared by developers and AI, with the architect owning the counsel given.')

h2('6B.  Roles Within a Larger Implementation')
body('This offering exists only when a build is underway. It is engaged as a defined role inside an implementation project rather than sold as a standalone engagement.')

h2('6.5  Architectural Decisions and Risk Ownership during the Build')
body('First-pass review and assessment are increasingly done by developers and AI. The architect’s value in delivery is not running those checks but owning the trade-off decisions they surface and being accountable for the risk that ships.')
bullet('Owning the architectural trade-offs — reliability, performance, cost, security, and speed-to-market are competing goods, and the right balance depends on this client’s context and risk appetite.')
bullet('Adjudicating the findings of AI- and developer-led reviews and assurance, deciding what must be fixed, what is acceptable, and what risk is knowingly accepted.')
bullet('Setting the architecture, standards, and guardrails up front that the build is then held against.')
bullet('Standing accountable for what is released — the decision and the risk acceptance, not the detection.')
bold_label('Engagement across the implementation lifecycle')
body('The role is heaviest at the start and at key decision gates, not uniformly throughout.')
bold_label('Project setup — high involvement — ', 'establish the architecture, standards, guardrails, and development approach the build will follow.')
bold_label('Throughout delivery — light, advisory involvement — ', 'available for design questions and to assess the impact of changing requirements.')
bold_label('Key gates — concentrated involvement — ', 'own the trade-off and risk decisions on major deliverables and design changes before they are committed.')
italic_note('Also applies as a standalone engagement:  as an independent architecture or assurance review commissioned outside a delivery role.')
bold_label('How and when this is engaged')
bold_label('Trigger — ', 'An implementation under way where consequential architectural trade-offs and release-risk decisions must be owned.')
bold_label('Engagement shape — ', 'A defined role within a delivery programme — front-loaded at setup, then concentrated at decision gates rather than continuous.')
bold_label('Staffing — ', 'Architect owns decisions and risk acceptance; first-pass review, assurance, and construction carried by developers and AI.')

# ─── 7. SUPPORTING DISCIPLINE ────────────────────────────────────────────────
h1('7. Supporting Discipline: Building for Knowledge')
body('The advisory work described in the preceding sections is only as credible as the judgment behind it. That judgment is kept honest by a deliberate, selective building practice — one that is not sold to clients and is not designed to compete on construction speed or volume, but is maintained because advisory credibility that is not tested against technical reality eventually drifts from it.')

h2('7.1  Activities')
bullet('Building working prototypes and proofs-of-concept in the AI-native stack — agentic workflows, retrieval-augmented patterns, evaluation strategies, and model orchestration — to establish first-hand what performs versus what only demonstrates well.')
bullet('Maintaining sufficient currency with cloud AI services (e.g. AWS Bedrock, Azure AI Foundry) and the underlying data platform plumbing to distinguish substance from vendor or team hand-waving.')
bullet('Building selectively before advising, so that recommendations on a pattern are informed by direct exposure to its failure modes and sharp edges.')

h2('7.2  Operating Principles')
bullet('This work is judged by whether it keeps the architect close to ground truth, not by whether it is billable.')
bullet('Specific tooling knowledge has a short half-life; the objective is grounding and the ability to detect what is real, not exhaustive mastery of current tools.')
bullet('If hands-on practice lapses, advisory judgment silently drifts from reality. Maintaining this layer is therefore a non-negotiable input to the credibility of the external layer.')

# ─── 8. THE GOVERNING FILTER ─────────────────────────────────────────────────
h1('8. The Governing Filter')
body('Because both construction and advisory work can be performed competently, the deciding question for any engagement is not which type of work it is, but whether it compounds:')
p = doc.add_paragraph()
p.style = 'Normal'
r = p.add_run('Does performing this work well make the architect harder to replace and more trusted next year — or does it simply produce one more competent output that anyone with current tooling could have produced?')
r.italic = True
body('Internal building passes this filter only when it demonstrably sharpens the external value layer. External engagements are prioritised where they make the architect harder to replace and more trusted, rather than where they merely produce volume.')

# ─── 9. SUMMARY ──────────────────────────────────────────────────────────────
h1('9. Summary')
body('The architect in an AI-native world places judgment, decision guidance, and accountability on the invoice, while retaining a deliberate, selective building practice as the credibility floor that keeps that judgment honest. Construction is no longer the headline offering; it is the root system beneath it. The role is defined not by a choice between building and advising, but by an operating model in which directed, accountable judgment is the value, and hands-on practice is the discipline that keeps the value real.')

doc.save(OUTPUT)
print('Saved:', OUTPUT)
