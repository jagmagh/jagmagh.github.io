from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
import docx

SOURCE = r'D:\www.jagmagh.com\jagmagh.github.io\inputs\Solution Architect Role Re-Definition.docx'
doc = Document(SOURCE)

print("=== DOCUMENT DEFAULT STYLES ===")
print(f"Default font: {doc.styles['Normal'].font.name}")
print(f"Default size: {doc.styles['Normal'].font.size}")

print("\n=== NAMED STYLES IN USE ===")
used = set()
for p in doc.paragraphs:
    if p.style: used.add(p.style.name)
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                if p.style: used.add(p.style.name)
for s in sorted(used):
    try:
        style = doc.styles[s]
    except KeyError:
        print(f"\nStyle: '{s}' [not found in styles collection]")
        continue
    f = style.font
    pf = style.paragraph_format
    print(f"\nStyle: '{s}'")
    print(f"  font.name={f.name}, size={f.size}, bold={f.bold}, italic={f.italic}, color={f.color.rgb if f.color and f.color.type else 'inherited'}")
    try:
        print(f"  space_before={pf.space_before}, space_after={pf.space_after}, left_indent={pf.left_indent}")
    except:
        pass

print("\n=== PARAGRAPH-BY-PARAGRAPH DUMP ===")
for i, p in enumerate(doc.paragraphs):
    if not p.text.strip():
        continue
    pf = p.paragraph_format
    print(f"\n[{i}] style='{p.style.name}' | alignment={p.alignment}")
    print(f"     space_before={pf.space_before} space_after={pf.space_after} left_indent={pf.left_indent}")
    print(f"     text (first 80): {p.text[:80]!r}")
    for j, run in enumerate(p.runs):
        f = run.font
        color = None
        try:
            color = f.color.rgb
        except:
            pass
        highlight = run.font.highlight_color
        shade = None
        # Check run shading
        rpr = run._r.find(qn('w:rPr'))
        if rpr is not None:
            shd = rpr.find(qn('w:shd'))
            if shd is not None:
                shade = shd.get(qn('w:fill'))
        print(f"     run[{j}]: bold={f.bold} italic={f.italic} size={f.size} name={f.name} color={color} highlight={highlight} shade={shade} | {run.text[:60]!r}")

print("\n=== TABLE DUMP ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti}: style='{table.style.name}', cols={len(table.columns)}, rows={len(table.rows)}")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            # Cell background
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            bg = None
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    bg = shd.get(qn('w:fill'))
            for p in cell.paragraphs:
                if p.text.strip():
                    pf = p.paragraph_format
                    runs_info = []
                    for r in p.runs:
                        f = r.font
                        c = None
                        try: c = f.color.rgb
                        except: pass
                        runs_info.append(f"bold={f.bold},size={f.size},color={c},italic={f.italic}")
                    print(f"  [{ri},{ci}] bg={bg} | '{p.text[:50]}' | runs: {runs_info}")

print("\n=== SECTION / PAGE SETUP ===")
for s in doc.sections:
    print(f"  margins: top={s.top_margin}, bottom={s.bottom_margin}, left={s.left_margin}, right={s.right_margin}")
    print(f"  page: width={s.page_width}, height={s.page_height}")
