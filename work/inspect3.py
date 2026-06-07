from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SOURCE = r'D:\www.jagmagh.com\jagmagh.github.io\inputs\Solution Architect Role Re-Definition.docx'
doc = Document(SOURCE)

# ── Governing filter question full XML ────────────────────────────────────────
print("=== GOVERNING FILTER PARAGRAPH XML ===")
for i, p in enumerate(doc.paragraphs):
    if 'Does performing this work well' in p.text:
        print(f"Paragraph index: {i}")
        print(etree.tostring(p._p, pretty_print=True).decode())

# ── Landscape / section breaks in document ────────────────────────────────────
print("\n=== ALL INLINE sectPr (section breaks) ===")
for i, p in enumerate(doc.paragraphs):
    pPr = p._p.find(qn('w:pPr'))
    if pPr is not None:
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            print(f"\nParagraph {i}: '{p.text[:60]}'")
            print(etree.tostring(sectPr, pretty_print=True).decode())

print("\n=== DOCUMENT BODY sectPr ===")
body = doc.element.body
sectPr = body.find(qn('w:sectPr'))
if sectPr is not None:
    print(etree.tostring(sectPr, pretty_print=True).decode())

# ── "8. The Governing Filter" heading paragraph ───────────────────────────────
print("\n=== GOVERNING FILTER SECTION HEADING ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('8.') or p.text.strip() == 'The Governing Filter' or '8. The Governing Filter' in p.text:
        print(f"[{i}] '{p.text}'")
        pf = p.paragraph_format
        print(f"  style={p.style.name if p.style else 'None'}, before={pf.space_before}, after={pf.space_after}")
        for r in p.runs:
            f = r.font
            c = None
            try: c = f.color.rgb
            except: pass
            print(f"  run: b={f.bold}, i={f.italic}, sz={f.size}, c={c} | '{r.text}'")
