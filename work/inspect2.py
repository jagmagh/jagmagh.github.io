from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SOURCE = r'D:\www.jagmagh.com\jagmagh.github.io\inputs\Solution Architect Role Re-Definition.docx'
doc = Document(SOURCE)

# ── 1. Heading 1 style from XML ──────────────────────────────────────────────
print("=== HEADING 1 STYLE XML ===")
for style in doc.styles.element.findall(qn('w:style')):
    sid = style.get(qn('w:styleId'))
    if sid and 'Heading1' in sid:
        print(etree.tostring(style, pretty_print=True).decode())

# ── 2. Full table dump (cell backgrounds + run formatting) ───────────────────
print("\n=== TABLE FULL DUMP ===")
for ti, table in enumerate(doc.tables):
    sname = table.style.name if table.style else 'None'
    print(f"\nTable {ti}: style='{sname}', rows={len(table.rows)}, cols={len(table.columns)}")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            bg = 'none'
            border_info = ''
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    bg = shd.get(qn('w:fill'), 'none')
                    bg_color = shd.get(qn('w:color'), '')
                    bg_val = shd.get(qn('w:val'), '')
                    bg = f"fill={bg} color={bg_color} val={bg_val}"
            for p in cell.paragraphs:
                if not p.text.strip():
                    continue
                runs_summary = []
                for r in p.runs:
                    f = r.font
                    c = None
                    try: c = f.color.rgb
                    except: pass
                    runs_summary.append(f"b={f.bold},sz={f.size},c={c},i={f.italic}|'{r.text[:30]}'")
                print(f"  [{ri},{ci}] bg={bg}")
                print(f"         text='{p.text[:60]}'")
                print(f"         runs={runs_summary}")

# ── 3. Key paragraph patterns (indented notes, "How and when" blocks) ────────
print("\n=== INDENTED / SPECIAL PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    pf = p.paragraph_format
    li = pf.left_indent
    if li and li > 0 and p.text.strip():
        runs_summary = []
        for r in p.runs:
            f = r.font
            c = None
            try: c = f.color.rgb
            except: pass
            runs_summary.append(f"b={f.bold},i={f.italic},sz={f.size},c={c}|'{r.text[:40]}'")
        print(f"[{i}] indent={li} before={pf.space_before} after={pf.space_after}")
        print(f"     '{p.text[:70]}'")
        print(f"     {runs_summary}")

# ── 4. "6A" and "6B" section heading paragraphs ───────────────────────────────
print("\n=== 6A/6B SECTION HEADING PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('6A') or p.text.strip().startswith('6B'):
        pf = p.paragraph_format
        runs_summary = []
        for r in p.runs:
            f = r.font
            c = None
            try: c = f.color.rgb
            except: pass
            runs_summary.append(f"b={f.bold},i={f.italic},sz={f.size},c={c}|'{r.text[:40]}'")
        print(f"[{i}] style='{p.style.name}' before={pf.space_before} after={pf.space_after} indent={pf.left_indent}")
        print(f"     '{p.text[:80]}'")
        print(f"     {runs_summary}")

# ── 5. "Carries into delivery" / "Also applies" pattern ─────────────────────
print("\n=== CARRIES/ALSO APPLIES PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('Carries into') or t.startswith('Also applies'):
        pf = p.paragraph_format
        runs_summary = []
        for r in p.runs:
            f = r.font
            c = None
            try: c = f.color.rgb
            except: pass
            runs_summary.append(f"b={f.bold},i={f.italic},sz={f.size},c={c}|'{r.text[:40]}'")
        print(f"[{i}] style='{p.style.name}' before={pf.space_before} after={pf.space_after} indent={pf.left_indent}")
        print(f"     '{p.text[:80]}'")
        print(f"     {runs_summary}")
